"""
Markdown 导出模块
支持导出为 HTML、PDF、Word 格式
"""

import os
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
import markdown


class MarkdownExporter:
    """Markdown 导出器"""
    
    def __init__(self, markdown_content: str):
        """
        初始化导出器
        
        Args:
            markdown_content: Markdown 格式的内容
        """
        self.markdown_content = markdown_content
        self.html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
    
    def export_html(self, output_path: str, title: str = "Markdown Document") -> str:
        """
        导出为 HTML 文件
        
        Args:
            output_path: 输出文件路径
            title: 文档标题
            
        Returns:
            输出文件路径
        """
        html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', Courier, monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        pre code {{
            padding: 0;
            background: none;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 15px;
            color: #666;
            margin-left: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 15px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f4f4f4;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 24px;
            margin-bottom: 16px;
        }}
    </style>
</head>
<body>
    {self.html_content}
</body>
</html>"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_template)
        
        return output_path
    
    def export_pdf(self, output_path: str, title: str = "Markdown Document") -> str:
        """
        导出为 PDF 文件
        
        Args:
            output_path: 输出文件路径
            title: 文档标题
            
        Returns:
            输出文件路径
        """
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # 自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=TA_LEFT
        )
        
        # 添加标题
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # 将 HTML 转换为 PDF 内容
        # 简化处理：将 HTML 标签移除，保留基本格式
        content = self._html_to_text(self.html_content)
        
        for line in content.split('\n'):
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
        
        doc.build(story)
        return output_path
    
    def _html_to_text(self, html_content: str) -> str:
        """
        将 HTML 内容转换为纯文本（简化版本）
        
        Args:
            html_content: HTML 格式内容
            
        Returns:
            纯文本内容
        """
        import re
        
        # 移除 HTML 标签
        text = re.sub(r'<[^>]+>', '', html_content)
        
        # 解码 HTML 实体
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&amp;', '&')
        text = text.replace('&quot;', '"')
        
        return text
    
    def export_word(self, output_path: str, title: str = "Markdown Document") -> str:
        """
        导出为 Word 文件
        
        Args:
            output_path: 输出文件路径
            title: 文档标题
            
        Returns:
            输出文件路径
        """
        doc = Document()
        
        # 添加标题
        doc.add_heading(title, 0)
        
        # 添加内容（简化处理）
        content = self._html_to_text(self.html_content)
        
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_path)
        return output_path


def export_markdown(content: str, output_path: str, format: str = 'html', title: str = "Document") -> str:
    """
    导出 Markdown 内容的便捷函数
    
    Args:
        content: Markdown 内容
        output_path: 输出路径
        format: 导出格式 ('html', 'pdf', 'word')
        title: 文档标题
        
    Returns:
        输出文件路径
    """
    exporter = MarkdownExporter(content)
    
    if format.lower() == 'html':
        return exporter.export_html(output_path, title)
    elif format.lower() == 'pdf':
        return exporter.export_pdf(output_path, title)
    elif format.lower() == 'word':
        return exporter.export_word(output_path, title)
    else:
        raise ValueError(f"不支持的格式：{format}")
